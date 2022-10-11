from docx import Document
from docx.shared import Inches

document = Document()

# prfile picture

document.add_picture(
    'pexels-allan-mas-5622388.jpg', width=Inches(2.0))

# name phone number and email details
name = input('what is your name? ')
phone_number = input('what is your phone_number? ')
email = input('what is your email? ')

document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email)

# about me
document.add_heading('About me')
document.add_paragraph(
    input('Tell me about yourself? ')
)

# work experience
document.add_heading('work experience')
p = document.add_paragraph()

company = input('Enter company ')
from_date = input('From Date')
to_date = input('To Date ')

p.add_run(company + ' ').bold = True
p.add_run(from_date  + '-' + to_date + '\n').italic = True

experience_details = input(
    'Describe your experience at ' + company)
p.add_run(experience_details)


document.save('cv.docx')










 



  


    


    



